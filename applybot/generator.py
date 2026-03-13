"""Generate tailored resume and cover letter .docx files."""
from __future__ import annotations
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from applybot.scraper import JobPost


def make_filename(company: str, role: str, doc_type: str) -> str:
    slug = re.sub(r"[^\w]", "_", company)[:20]
    role_slug = re.sub(r"[^\w]", "_", role)[:25]
    return f"{slug}_{role_slug}_{doc_type.capitalize()}.docx"


def _parse_markdown_resume(md_text: str) -> dict[str, Any]:
    """Very lightweight markdown → structured dict. No external deps."""
    lines = md_text.splitlines()
    result: dict[str, Any] = {"name": "", "contact": "", "sections": []}
    current_section: dict[str, Any] | None = None

    for line in lines:
        if line.startswith("# "):
            result["name"] = line[2:].strip()
        elif line.startswith("## "):
            if current_section:
                result["sections"].append(current_section)
            current_section = {"heading": line[3:].strip(), "content": []}
        elif current_section is not None:
            current_section["content"].append(line)

    if current_section:
        result["sections"].append(current_section)

    return result


def generate_resume(
    job: JobPost,
    resume_text: str,
    config: dict[str, Any],
    output_dir: Path,
) -> Path:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    parsed = _parse_markdown_resume(resume_text)
    doc = Document()

    # Name heading
    heading = doc.add_heading(parsed["name"] or config["name"], level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact line
    contact_parts = [config["email"], config["phone"]]
    if config.get("linkedin_url"):
        contact_parts.append(config["linkedin_url"])
    if config.get("portfolio_url"):
        contact_parts.append(config["portfolio_url"])
    contact_para = doc.add_paragraph(" | ".join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sections
    for section in parsed["sections"]:
        doc.add_heading(section["heading"], level=2)
        for line in section["content"]:
            if not line.strip():
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line.strip())

    out_path = output_dir / make_filename(job.company, job.title, "resume")
    doc.save(str(out_path))
    return out_path


def generate_cover_letter(
    job: JobPost,
    resume_text: str,
    config: dict[str, Any],
    output_dir: Path,
) -> Path:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    name = config["name"]
    company = job.company
    role = job.title

    doc.add_heading(f"{name}", level=1)
    doc.add_paragraph(
        f"{config['email']} | {config['phone']} | {config.get('linkedin_url', '')}"
    )
    doc.add_paragraph("")

    doc.add_paragraph(f"Hiring Team at {company}")
    doc.add_paragraph("")

    body = (
        f"Dear Hiring Team,\n\n"
        f"I'm excited to apply for the {role} position at {company}. "
        f"With {config.get('years_of_experience', 'several')} years of experience, "
        f"I bring a track record of delivering results in fast-paced environments.\n\n"
        f"I was drawn to {company} because of your focus on building impactful products. "
        f"My background aligns well with what you're looking for in this role, and I'm "
        f"confident I can make a meaningful contribution from day one.\n\n"
        f"I'd love the opportunity to discuss how my experience can benefit your team. "
        f"Thank you for your consideration.\n\n"
        f"Best regards,\n{name}"
    )

    for para_text in body.split("\n\n"):
        doc.add_paragraph(para_text)

    out_path = output_dir / make_filename(job.company, job.title, "coverletter")
    doc.save(str(out_path))
    return out_path
