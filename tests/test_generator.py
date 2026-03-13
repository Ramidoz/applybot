import pytest
from pathlib import Path
from docx import Document
from applybot.generator import generate_resume, generate_cover_letter, make_filename


def test_make_filename_sanitizes():
    name = make_filename("Acme Corp", "Senior Software Engineer", "resume")
    assert " " not in name
    assert name.endswith(".docx")
    assert "Acme" in name


def test_generate_resume_creates_docx(tmp_project, sample_config):
    from applybot.scraper import JobPost
    job = JobPost(
        title="Senior Software Engineer",
        company="Acme",
        url="https://acme.com/jobs/1",
        description="Python FastAPI PostgreSQL Docker",
        platform="linkedin",
    )
    out_path = generate_resume(
        job=job,
        resume_text=(tmp_project / "resume.md").read_text(),
        config=sample_config,
        output_dir=tmp_project / "output" / "resumes",
    )
    assert out_path.exists()
    assert out_path.suffix == ".docx"
    doc = Document(str(out_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane Smith" in full_text


def test_generate_cover_letter_creates_docx(tmp_project, sample_config):
    from applybot.scraper import JobPost
    job = JobPost(
        title="Backend Engineer",
        company="BetaCo",
        url="https://betaco.com/jobs/2",
        description="We need a backend engineer with Python skills.",
        platform="indeed",
    )
    out_path = generate_cover_letter(
        job=job,
        resume_text=(tmp_project / "resume.md").read_text(),
        config=sample_config,
        output_dir=tmp_project / "output" / "cover_letters",
    )
    assert out_path.exists()
    doc = Document(str(out_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "BetaCo" in full_text
    assert "Jane Smith" in full_text
